import os
import uuid
import shutil
import requests
from flask import Flask, request, jsonify, render_template, session
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from PyPDF2 import PdfReader
from docx import Document
import logging
import openai
from dotenv import load_dotenv
from duckduckgo_search import DDGS
import boto3

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY")  # e.g., "supersecretkey"

# Logging setup
logging.basicConfig(filename="app.log", level=logging.INFO)

# OpenRouter DeepSeek R1 API Settings (using OpenAI's compatible interface)
openrouter_api_key = os.getenv("DEEPSEEK_API_KEY")
openai.api_base = "https://openrouter.ai/api/v1"
openai.api_key = openrouter_api_key
DEEPEEK_MODEL = "deepseek/deepseek-r1:free"

# S3 Bucket name from environment
S3_BUCKET = os.getenv("S3_BUCKET")

# Text Splitter for breaking documents into chunks
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

#########################################
# Helper: Sync a local directory to S3
#########################################
def sync_directory_to_s3(local_directory, bucket_name, s3_prefix):
    s3_client = boto3.client('s3')
    for root, dirs, files in os.walk(local_directory):
        for file in files:
            local_path = os.path.join(root, file)
            # Compute the relative path so we can mirror directory structure in S3
            relative_path = os.path.relpath(local_path, local_directory)
            s3_path = os.path.join(s3_prefix, relative_path)
            try:
                s3_client.upload_file(local_path, bucket_name, s3_path)
                logging.info(f"Synced {local_path} to s3://{bucket_name}/{s3_path}")
            except Exception as e:
                logging.error(f"Failed to sync {local_path} to s3://{bucket_name}/{s3_path}: {str(e)}")

#########################################
# Global Vector Store (Preloaded Documents)
#########################################
def load_global_vector_db():
    # Use a local directory that we later sync to S3
    global_db_path = "global_chroma"
    os.makedirs(global_db_path, exist_ok=True)
    return Chroma(
        persist_directory=global_db_path,
        embedding_function=HuggingFaceEmbeddings(model_name="paraphrase-MiniLM-L6-v2")
    )

global_vector_db = load_global_vector_db()

def extract_text_from_path(file_path):
    try:
        if file_path.endswith(".pdf"):
            with open(file_path, "rb") as f:
                reader = PdfReader(f)
                return "".join(page.extract_text() for page in reader.pages if page.extract_text())
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs)
        elif file_path.endswith(".txt"):
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            return None
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {str(e)}")
        return None

def load_global_documents():
    global_docs_dir = "global_docs"
    if os.path.exists(global_docs_dir):
        for filename in os.listdir(global_docs_dir):
            file_path = os.path.join(global_docs_dir, filename)
            text = extract_text_from_path(file_path)
            if not text:
                continue
            chunks = text_splitter.split_text(text)
            documents = [{"content": chunk, "metadata": {"source": filename}} for chunk in chunks]
            global_vector_db.add_texts(
                [doc["content"] for doc in documents],
                metadatas=[doc["metadata"] for doc in documents]
            )
            logging.info(f"Preloaded {filename} with {len(chunks)} chunks.")
    # After loading, sync the global vector store to S3
    sync_directory_to_s3("global_chroma", S3_BUCKET, "global_chroma")

# Preload global documents at startup
load_global_documents()

#########################################
# Session-specific (Temporary) Store
#########################################
def get_session_vector_db():
    if "session_id" not in session:
        session["session_id"] = str(uuid.uuid4())
    session_db_path = f"temp_chroma/{session['session_id']}"
    os.makedirs(session_db_path, exist_ok=True)
    return Chroma(
        persist_directory=session_db_path,
        embedding_function=HuggingFaceEmbeddings(model_name="paraphrase-MiniLM-L6-v2")
    )

#########################################
# File Processing Functions
#########################################
def extract_text_from_file(file):
    try:
        if file.filename.endswith(".pdf"):
            reader = PdfReader(file)
            return "".join(page.extract_text() for page in reader.pages if page.extract_text())
        elif file.filename.endswith(".docx"):
            doc = Document(file)
            return "\n".join(para.text for para in doc.paragraphs)
        elif file.filename.endswith(".txt"):
            return file.read().decode("utf-8")
        else:
            raise Exception("Unsupported file format")
    except Exception as e:
        logging.error(f"Error extracting text from file {file.filename}: {str(e)}")
        return None

#########################################
# Routes
#########################################

# Upload Document Route (for session-specific temporary uploads)
@app.route("/upload", methods=["POST"])
def upload_document():
    file = request.files.get("file")
    if not file:
        logging.error("No file uploaded")
        return jsonify({"error": "No file uploaded"}), 400

    # Check if file is already preloaded (exists in global_docs)
    if file.filename in os.listdir("global_docs"):
        return jsonify({"message": "This document is already preloaded and available for queries!"}), 200

    text = extract_text_from_file(file)
    if not text:
        logging.error(f"Failed to extract text from {file.filename}")
        return jsonify({"error": "Failed to extract text from the file"}), 400

    chunks = text_splitter.split_text(text)
    documents = [{"content": chunk, "metadata": {"source": file.filename}} for chunk in chunks]

    session_db = get_session_vector_db()
    session_db.add_texts(
        [doc["content"] for doc in documents],
        metadatas=[doc["metadata"] for doc in documents]
    )
    logging.info(f"Uploaded {file.filename} with {len(chunks)} chunks for session {session['session_id']}.")

    # Sync session-specific vector store to S3
    sync_directory_to_s3(f"temp_chroma/{session['session_id']}", S3_BUCKET, f"temp_chroma/{session['session_id']}")

    return jsonify({"message": "Document uploaded successfully!"}), 200

# Fallback Internet Search using DuckDuckGo
def internet_search(query):
    try:
        with DDGS() as ddgs:
            results = ddgs.text(query, max_results=3)
            if results:
                return "\n".join(item.get("body", item.get("snippet", "")) for item in results if item.get("body") or item.get("snippet"))
            else:
                return "No relevant results found."
    except Exception as e:
        logging.error(f"Internet search error: {str(e)}")
        return "Error occurred while searching the internet."

# Query DeepSeek R1 via OpenRouter (using OpenAI ChatCompletion interface)
def query_deepseek(prompt):
    try:
        response = openai.ChatCompletion.create(
            model=DEEPEEK_MODEL,
            messages=[{"role": "user", "content": prompt}]
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logging.error(f"Error querying DeepSeek R1: {str(e)}")
        return f"Error: {str(e)}"

# Ask Question Route (combining global and session context)
@app.route("/ask", methods=["POST"])
def ask_question():
    data = request.get_json()
    question = data.get("question", "")
    if not question:
        return jsonify({"error": "No question provided"}), 400

    # Retrieve context from session-specific vector store
    session_vector_db = get_session_vector_db()
    session_docs = session_vector_db.similarity_search(question, k=3)
    session_context = "\n".join([doc.page_content for doc in session_docs])

    # Retrieve context from global vector store
    global_docs = global_vector_db.similarity_search(question, k=3)
    global_context = "\n".join([doc.page_content for doc in global_docs])

    combined_context = (session_context + "\n" + global_context).strip()

    # Fallback to internet search if no context found
    if not combined_context:
        combined_context = internet_search(question)
        logging.info("Using internet search fallback for context.")
#         Answer the question based solely on the provided context.
# If you don't know the answer, say you don't know.

    prompt = f"""Check with the document provided first. If not found in given document, search the internet. and also specify that this is general information form the internet.
    

Context:
{combined_context}

Question: {question}
Answer:"""

    answer = query_deepseek(prompt)
    logging.info(f"Question: {question} | Answer: {answer}")
    return jsonify({"answer": answer})

# Clear Session and Temporary Data Route
@app.route("/end_session", methods=["POST"])
def end_session():
    if "session_id" in session:
        session_db_path = f"temp_chroma/{session['session_id']}"
        if os.path.exists(session_db_path):
            shutil.rmtree(session_db_path)
        session.clear()
        response = jsonify({"message": "Session cleared successfully."})
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        return response, 200
    return jsonify({"message": "No active session to clear."}), 400

# Home Page Route with Sidebar listing preloaded documents
@app.route("/")
def home():
    preloaded_files = os.listdir("global_docs") if os.path.exists("global_docs") else []
    return render_template("index.html", preloaded_files=preloaded_files)

if __name__ == "__main__":
    app.run(debug=True)